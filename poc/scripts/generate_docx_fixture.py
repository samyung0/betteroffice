#!/usr/bin/env python3
"""Build the deterministic DOCX used by the browser round-trip proof."""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ACCENT = "2563EB"
INK = RGBColor(15, 23, 42)
MUTED = RGBColor(71, 85, 105)
MARKER = "EVO_EDIT_MARKER_DOCX"


def set_cell_fill(cell, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    properties.append(shading)


def add_hyperlink(paragraph, label: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), ACCENT)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend([color, underline])
    text = OxmlElement("w:t")
    text.text = label
    run.extend([properties, text])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.size = Pt(9)
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.extend([field_begin, instruction, field_end])


def build(output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.core_properties.title = "BetterOffice browser round-trip fixture"
    document.core_properties.subject = "Known OOXML features for loss detection"
    document.core_properties.keywords = "betteroffice, ooxml, roundtrip, fixture"

    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(7)

    for style_name, size in (("Title", 30), ("Heading 1", 19), ("Heading 2", 13)):
        style = document.styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = INK

    header = section.header.paragraphs[0]
    header.text = "EVO OFFICE · ROUND-TRIP PROOF"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.bold = True
    header.runs[0].font.color.rgb = MUTED
    add_page_field(section.footer.paragraphs[0])

    title = document.add_paragraph(style="Title")
    title.add_run("Browser document fixture")
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    subtitle_run = subtitle.add_run(
        "A deliberately compact document with enough structure to expose lossy open/edit/save behavior."
    )
    subtitle_run.font.size = Pt(13)
    subtitle_run.font.color.rgb = MUTED

    marker = document.add_paragraph()
    marker.paragraph_format.space_after = Pt(14)
    marker_run = marker.add_run(MARKER)
    marker_run.bold = True
    marker_run.font.color.rgb = RGBColor(255, 255, 255)
    marker_run.font.size = Pt(11)
    marker_xml = marker._p.get_or_add_pPr()
    marker_shading = OxmlElement("w:shd")
    marker_shading.set(qn("w:fill"), ACCENT)
    marker_xml.append(marker_shading)

    document.add_heading("Formatting and relationships", level=1)
    paragraph = document.add_paragraph()
    paragraph.add_run("This paragraph mixes ")
    paragraph.add_run("bold").bold = True
    paragraph.add_run(", ")
    paragraph.add_run("italic").italic = True
    paragraph.add_run(", and a real external ")
    add_hyperlink(paragraph, "hyperlink relationship", "https://betteroffice.dev")
    paragraph.add_run(". The link target should survive a no-op and edited save.")

    for item in (
        "Keep paragraph and run formatting intact.",
        "Keep numbered-list structure and indentation intact.",
        "Keep tables, merged cells, headers, footers, and section settings intact.",
    ):
        document.add_paragraph(item, style="List Number")

    document.add_heading("Structured content", level=1)
    table = document.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    widths = (2.45, 1.25, 1.35, 1.75)
    headers = ("Capability", "Viewer", "Editor", "Expected result")
    for index, (cell, label, width) in enumerate(zip(table.rows[0].cells, headers, widths)):
        cell.width = Inches(width)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_fill(cell, ACCENT)
        run = cell.paragraphs[0].add_run(label)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        if index > 0:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    rows = (
        ("Open OOXML", "Yes", "Yes", "No parse error"),
        ("No-op save", "—", "Yes", "Reopens cleanly"),
        ("Targeted text edit", "—", "Yes", "Marker changes only"),
    )
    for row_values in rows:
        cells = table.add_row().cells
        for cell, value, width in zip(cells, row_values, widths):
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = value
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    note_row = table.add_row().cells
    merged = note_row[0].merge(note_row[3])
    set_cell_fill(merged, "EFF6FF")
    merged.text = "Merged-cell sentinel: EVO_TABLE_MERGE_SENTINEL"
    merged.paragraphs[0].runs[0].italic = True

    document.add_page_break()
    document.add_heading("Second-page preservation sentinel", level=1)
    document.add_paragraph(
        "This page proves that pagination, header/footer references, and content after a hard page break remain addressable."
    )
    document.add_heading("Non-ASCII text", level=2)
    document.add_paragraph("Résumé · naïve · Ελληνικά · 中文 · العربية · emoji ✓")

    second_section = document.add_section(WD_SECTION.NEW_PAGE)
    second_section.orientation = section.orientation
    second_section.header.is_linked_to_previous = True
    second_section.footer.is_linked_to_previous = True
    document.add_heading("Section boundary", level=1)
    document.add_paragraph(
        "EVO_SECTION_SENTINEL — a second OOXML section should survive both save paths."
    )

    document.save(output_path)


if __name__ == "__main__":
    if len(sys.argv) != 2:
        raise SystemExit("usage: generate_docx_fixture.py OUTPUT.docx")
    build(Path(sys.argv[1]).resolve())
